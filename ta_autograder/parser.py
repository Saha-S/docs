"""
parser.py – Extract plain text from student submission files.

Supported formats:
  - PDF  : pdfplumber (primary) → PyMuPDF/fitz (fallback)
  - DOCX : python-docx
"""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_text(file_path: Path) -> str:
    """Return the full text content of *file_path*.

    Raises
    ------
    ValueError
        If the file extension is not supported.
    RuntimeError
        If all parsing attempts fail.
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(file_path)
    if suffix == ".docx":
        return _extract_docx(file_path)

    raise ValueError(f"Unsupported file format: '{suffix}' ({file_path.name})")


# ---------------------------------------------------------------------------
# PDF helpers
# ---------------------------------------------------------------------------

def _extract_pdf(path: Path) -> str:
    """Try pdfplumber first, fall back to PyMuPDF."""
    text = _pdf_pdfplumber(path)
    if text and text.strip():
        return text

    logger.warning(
        "pdfplumber returned empty text for %s – trying PyMuPDF fallback", path.name
    )
    text = _pdf_pymupdf(path)
    if text and text.strip():
        return text

    raise RuntimeError(f"Could not extract text from PDF: {path}")


def _pdf_pdfplumber(path: Path) -> str:
    try:
        import pdfplumber  # type: ignore

        pages: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
        return "\n".join(pages)
    except Exception as exc:  # noqa: BLE001
        logger.debug("pdfplumber failed for %s: %s", path.name, exc)
        return ""


def _pdf_pymupdf(path: Path) -> str:
    try:
        import fitz  # type: ignore  # PyMuPDF

        doc = fitz.open(str(path))
        pages: list[str] = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(pages)
    except Exception as exc:  # noqa: BLE001
        logger.debug("PyMuPDF failed for %s: %s", path.name, exc)
        return ""


# ---------------------------------------------------------------------------
# DOCX helper
# ---------------------------------------------------------------------------

def _extract_docx(path: Path) -> str:
    try:
        from docx import Document  # type: ignore

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also capture text inside tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)
        return "\n".join(paragraphs)
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"python-docx failed for {path}: {exc}") from exc
